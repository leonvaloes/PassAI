from __future__ import annotations

from pathlib import Path
import re
import sys
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = PROJECT_ROOT / "layoutCV" / "layout.docx"


def split_items(text: str) -> List[str]:
    normalized = text.replace(";", ",")
    return [item.strip() for item in normalized.split(",") if item.strip()]


def parse_markdown(markdown_path: Path) -> Dict[str, object]:
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    data: Dict[str, object] = {
        "nome": "",
        "cargo": "",
        "cidade": "",
        "estado": "",
        "email": "",
        "telefone": "",
        "linkedin": "",
        "github": "",
        "resumo": "",
        "competencias": [],
        "educacao": "Bacharelado em Sistemas de Informação - Unoeste\nJan/2022 - Dez/2026",
        "experiencias": [],
    }

    current_section = ""
    current_experience = None
    objective_lines: List[str] = []
    summary_lines: List[str] = []
    education_lines: List[str] = []
    skills: List[str] = []

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            data["nome"] = line[2:].strip()
            continue

        if line.startswith("## "):
            current_section = line[3:].strip().lower()
            current_experience = None
            continue

        if line.startswith("### ") and "experiência" in current_section:
            heading = line[4:].strip()
            if " - " in heading:
                company, role = heading.split(" - ", 1)
            else:
                company, role = heading, ""
            current_experience = {
                "empresa": company.strip(),
                "cargo": role.strip(),
                "periodo": "",
                "bullets": [],
                "_skip_continuation": False,
            }
            data["experiencias"].append(current_experience)
            continue

        if current_section == "":
            if line.startswith("LinkedIn:"):
                pieces = [part.strip() for part in line.split("|")]
                data["linkedin"] = pieces[0].replace("LinkedIn:", "").strip()
                if len(pieces) > 1:
                    data["github"] = pieces[1].replace("GitHub:", "").strip()
            elif "|" in line:
                parts = [part.strip() for part in line.split("|")]
                if parts:
                    data["email"] = parts[0]
                if len(parts) > 1:
                    data["telefone"] = parts[1]
            elif "," in line:
                city, state = [part.strip() for part in line.split(",", 1)]
                data["cidade"] = city
                data["estado"] = state
            continue

        if current_section == "objetivo":
            objective_lines.append(line)
            continue

        if current_section == "resumo profissional":
            summary_lines.append(line)
            continue

        if current_section == "competências técnicas":
            if line.startswith("- "):
                content = line[2:].strip()
                if ":" in content:
                    _, content = content.split(":", 1)
                for item in split_items(content):
                    skills.append(item)
            continue

        if current_section == "educação":
            education_lines.append(line)
            continue

        if "experiência" in current_section and current_experience:
            if not current_experience["periodo"]:
                current_experience["periodo"] = line
            elif line.startswith("- "):
                bullet = line[2:].strip()
                if not bullet.startswith("Tecnologias:"):
                    current_experience["bullets"].append(bullet)
                    current_experience["_skip_continuation"] = False
                else:
                    current_experience["_skip_continuation"] = True
            elif raw_line.startswith((" ", "\t")) and current_experience["bullets"]:
                if not current_experience.get("_skip_continuation"):
                    current_experience["bullets"][-1] = (
                        f"{current_experience['bullets'][-1]} {line}"
                    )

    objective = " ".join(objective_lines).strip()
    if objective:
        data["cargo"] = objective.split(",", 1)[0].strip()

    data["resumo"] = " ".join(summary_lines).strip()
    if education_lines:
        data["educacao"] = "\n".join(education_lines)

    deduped_skills = []
    seen = set()
    for skill in skills:
        normalized = skill.lower()
        if normalized not in seen and len(skill) <= 40:
            seen.add(normalized)
            deduped_skills.append(skill)
    data["competencias"] = deduped_skills[:18]

    return data


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def normalize_url(display_text: str, kind: str) -> str:
    text = display_text.strip()
    if text.startswith(("http://", "https://")):
        return text

    if kind == "linkedin":
        slug = text.replace("linkedin.com/in/", "").replace("linkedin.com/", "").strip("/")
        return f"https://www.linkedin.com/in/{slug}/"

    if kind == "github":
        slug = text.replace("github.com/", "").strip("/")
        return f"https://github.com/{slug}"

    return f"https://{text}"


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    new_run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_paragraph_hyperlink(paragraph, display_text: str, url: str) -> None:
    set_paragraph_text(paragraph, "")
    add_hyperlink(paragraph, url, display_text)


def add_like(document, template_paragraph, text: str, *, bold: bool | None = None) -> None:
    paragraph = document.add_paragraph(style=template_paragraph.style)
    paragraph.alignment = template_paragraph.alignment or WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    if template_paragraph.runs:
        source_run = template_paragraph.runs[0]
        run.bold = source_run.bold if bold is None else bold
        run.italic = source_run.italic
        run.font.name = source_run.font.name
        run.font.size = source_run.font.size
    elif bold is not None:
        run.bold = bold


def fill_template(data: Dict[str, object], docx_path: Path, template_path: Path = DEFAULT_TEMPLATE) -> None:
    document = Document(str(template_path))

    replacements = {
        "{{NOME}}": str(data["nome"]),
        "{{CARGO}}": str(data["cargo"]),
        "{{CIDADE}}": str(data["cidade"]),
        "{{ESTADO}}": str(data["estado"]),
        "{{EMAIL}}": str(data["email"]),
        "{{TELEFONE}}": str(data["telefone"]),
        "{{RESUMO}}": str(data["resumo"]),
        "{{COMPETENCIAS}}": "\n".join(data["competencias"]),
    }

    for paragraph in document.paragraphs:
        text = paragraph.text
        for placeholder, value in replacements.items():
            text = text.replace(placeholder, value)
        if text != paragraph.text:
            set_paragraph_text(paragraph, text)

    for paragraph in document.paragraphs:
        if paragraph.text == "{{LINKEDIN}}":
            linkedin = str(data["linkedin"])
            set_paragraph_hyperlink(
                paragraph,
                linkedin,
                normalize_url(linkedin, "linkedin"),
            )
        elif paragraph.text == "{{GITHUB}}":
            github = str(data["github"])
            set_paragraph_hyperlink(
                paragraph,
                github,
                normalize_url(github, "github"),
            )

    for paragraph in document.paragraphs:
        if paragraph.text.startswith("Sistemas de informação"):
            set_paragraph_text(paragraph, str(data["educacao"]))
            break

    # Remove sample experience paragraphs after the section title and append
    # generated content using the original template paragraph styles.
    exp_heading_idx = None
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip().upper() == "EXPERIÊNCIA PROFISSIONAL":
            exp_heading_idx = idx
            break

    if exp_heading_idx is not None:
        style_samples = document.paragraphs[exp_heading_idx + 1 : exp_heading_idx + 6]
        company_style = style_samples[0]
        role_style = style_samples[1]
        period_style = style_samples[2]
        bullet_style = style_samples[3]

        for paragraph in list(document.paragraphs[exp_heading_idx + 1 :]):
            remove_paragraph(paragraph)

        for index, experience in enumerate(data["experiencias"]):
            add_like(document, company_style, str(experience["empresa"]), bold=True)
            if index > 0:
                document.add_paragraph(style=company_style.style)
            add_like(document, role_style, str(experience["cargo"]), bold=True)
            add_like(document, period_style, str(experience["periodo"]), bold=True)
            for bullet in experience["bullets"]:
                add_like(document, bullet_style, f"•  {bullet}")

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(docx_path))


def convert(markdown_path: Path, docx_path: Path) -> None:
    data = parse_markdown(markdown_path)
    fill_template(data, docx_path)


def main() -> int:
    if len(sys.argv) != 3:
        print("usage: python scripts/markdown_cv_to_docx.py <input.md> <output.docx>")
        return 2

    convert(Path(sys.argv[1]), Path(sys.argv[2]))
    print(f"Generated {sys.argv[2]} using {DEFAULT_TEMPLATE}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
