from docx import Document
from docx.shared import Pt, Inches

def create_default_template():
    doc = Document()
    
    # Title
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Name
    p = doc.add_paragraph()
    p.alignment = 1 # Center
    run = p.add_run('{{NOME}}')
    run.bold = True
    run.font.size = Pt(24)
    
    # Contact Info
    p = doc.add_paragraph()
    p.alignment = 1
    p.add_run('{{EMAIL}} | {{TELEFONE}} | {{LINKEDIN}}')
    
    # Summary
    doc.add_heading('RESUMO PROFISSIONAL', level=1)
    doc.add_paragraph('{{RESUMO}}')
    
    # Experience
    doc.add_heading('EXPERIÊNCIA PROFISSIONAL', level=1)
    doc.add_paragraph('Use a loop here in implementation, but for now just a placeholder')
    
    # Skills
    doc.add_heading('HABILIDADES', level=1)
    doc.add_paragraph('{{COMPETENCIAS}}')
    
    # Education
    doc.add_heading('FORMAÇÃO ACADÊMICA', level=1)
    doc.add_paragraph('{{EDUCACAO}}')
    
    doc.save('backend/layoutCV/layout.docx')
    print("Template created successfully at backend/layoutCV/layout.docx")

if __name__ == "__main__":
    create_default_template()
