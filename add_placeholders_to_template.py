"""
Script to add placeholders to existing template while preserving formatting
"""
from docx import Document
from pathlib import Path

# Load existing template
template_path = Path('layoutCV/layout.docx')
doc = Document(template_path)

# Mapping: paragraph index -> placeholder
replacements = {
    0: '{{NOME}}',  # Nome completo
    1: '{{CARGO}}',  # Cargo/palavra-chave
    2: '{{CIDADE}} – {{ESTADO}}',  # Localização
    3: '{{EMAIL}}',  # Email
    4: 'Telefone: {{TELEFONE}}',  # Telefone
    5: '{{LINKEDIN}}',  # LinkedIn
    # Parágrafo 7 é vazio
    # Parágrafo 8 é "RESUMO" (manter)
    10: '{{RESUMO}}',  # Conteúdo do resumo
    # Parágrafo 12 é "HABILIDADES" (manter)
    14: '{{COMPETENCIAS}}',  # Lista de habilidades
}

print(f"Total paragraphs: {len(doc.paragraphs)}")
print("\nCurrent content (first 20 paragraphs):")
for i, p in enumerate(doc.paragraphs[:20]):
    print(f"{i}: {p.text[:80]}")

print("\n\nApplying replacements...")
for idx, placeholder in replacements.items():
    if idx < len(doc.paragraphs):
        para = doc.paragraphs[idx]
        # Preserve all formatting, only change text
        if para.runs:
            # Replace text in first run, clear others
            para.runs[0].text = placeholder
            for run in para.runs[1:]:
                run.text = ""
        else:
            # No runs, add one
            para.add_run(placeholder)
        print(f"  [{idx}] {placeholder}")

# Save modified template
output_path = Path('layoutCV/layout_modified.docx')
doc.save(output_path)
print(f"\n✅ Modified template saved to: {output_path}")
print("Review the file, then rename it to 'layout.docx' if correct.")
