from __future__ import annotations

import json
import os
import re
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

try:
    from openai import OpenAI
except ImportError:  # pragma: no cover - optional dependency in test environments
    OpenAI = None


PROJECT_ROOT = Path(__file__).resolve().parent.parent

KNOWN_KEYWORDS = [
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "node",
    "node.js",
    "next.js",
    "vue",
    "angular",
    "spring",
    "spring boot",
    "fastapi",
    "django",
    "flask",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "sql",
    "postgresql",
    "mysql",
    "mongodb",
    "redis",
    "rabbitmq",
    "kafka",
    "microservices",
    "api",
    "rest",
    "graphql",
    "ci/cd",
    "git",
    "terraform",
    "linux",
]

ROLE_PATTERNS = [
    r"\bdesenvolvedor(?:a)?(?:\s+full[-\s]?stack)?\b",
    r"\bsoftware engineer\b",
    r"\bfull[-\s]?stack\b",
    r"\bbackend\b",
    r"\bfrontend\b",
    r"\bengenheir[oa]\s+de\s+software\b",
    r"\bjava developer\b",
    r"\bnode(?:\.js|js)? developer\b",
]

INVALID_SKILL_MARKERS = {
    "habilidade x",
    "tecnologia x",
    "requisitos das vagas",
    "requisitos da vaga",
    "skills",
    "skill",
}

RAG_CONTEXT = """
- ATS prioriza aderencia literal a tecnologias, stacks, cloud, banco de dados e termos operacionais.
- Preserve empresas, cargos e periodos do curriculo original.
- Tecnologias da vaga devem ser distribuidas entre resumo, habilidades e experiencias de forma natural.
- Bullets precisam ser tecnicos, especificos e plausiveis, evitando genericidade.
- Priorize linguagem direta, vocabulario de engenharia de software e contexto real de produto/sistema.
- Se a vaga mencionar arquitetura, qualidade, design patterns, CI/CD, observabilidade ou cloud, inclua isso nas experiencias.
- Resumes de software engineers com melhor desempenho normalmente combinam verbos fortes, ownership claro, stack explicita e impacto mensuravel ou operacional.
- Para humanos e ATS, prefira bullets escaneaveis com acao, contexto, tecnologias, colaboracao e resultado no mesmo item.
""".strip()

TEXT_COLOR = RGBColor(0x11, 0x11, 0x11)
LINK_COLOR = RGBColor(0x2B, 0x5C, 0x9C)
FONT_REGULAR = "Open Sans"
FONT_ITALIC = "Nunito"


class ResumeService:
    def __init__(self, output_dir: str | Path, enable_llm: bool = True):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.enable_llm = enable_llm
        self.env_file = PROJECT_ROOT / ".env"
        self.openai_api_key = self._read_env_value("OPENAI_API_KEY")
        self.openai_model = self._read_env_value("PASSAI_OPENAI_MODEL") or self._read_env_value("OPENAI_MODEL")
        self.openai_model = self.openai_model or "gpt-4o-mini"
        self.llm_client = self._build_llm_client()

    def parse_job(self, content: str, input_type: str) -> dict:
        normalized = content.strip()
        lines = [line.strip(" -:\t") for line in normalized.splitlines() if line.strip()]
        lowered = normalized.lower()

        title = self._find_first_match(
            lines,
            [
                r"cargo[:\s]+(.+)",
                r"vaga[:\s]+(.+)",
                r"position[:\s]+(.+)",
                r"role[:\s]+(.+)",
            ],
        )
        if not title:
            title = self._infer_job_title(lines, normalized)
        if not title:
            title = "Vaga sem titulo informado"

        company = self._find_first_match(
            lines,
            [
                r"empresa[:\s]+(.+)",
                r"company[:\s]+(.+)",
            ],
        )
        if not company and len(lines) > 1 and len(lines[1]) < 80:
            company = lines[1]
        company = company or "Empresa nao informada"

        modality = None
        if "hibrido" in lowered:
            modality = "hibrido"
        elif "remoto" in lowered or "remote" in lowered:
            modality = "remoto"
        elif "presencial" in lowered or "onsite" in lowered:
            modality = "presencial"

        location = self._find_first_match(
            lines,
            [
                r"local[:\s]+(.+)",
                r"location[:\s]+(.+)",
            ],
        )

        keywords = self.extract_keywords(normalized)
        behaviors = self.extract_behavior_keywords(normalized)

        return {
            "input_type": input_type,
            "content": normalized,
            "cargo": title,
            "empresa": company,
            "ats_detectado": "legacy-prompt" if self.llm_client else "rule-based",
            "requisitos_tecnicos": keywords,
            "requisitos_comportamentais": behaviors,
            "local": location,
            "modalidade": modality,
        }

    def _infer_job_title(self, lines: list[str], content: str) -> str | None:
        for line in lines[:12]:
            if len(line) > 90:
                continue
            lowered = line.lower()
            if any(term in lowered for term in ["great place", "marca mais valiosa", "vem pro", "oportunidade para"]):
                continue
            for pattern in ROLE_PATTERNS:
                match = re.search(pattern, lowered, flags=re.IGNORECASE)
                if match:
                    return line.strip()
        lowered = content.lower()
        if "full-stack" in lowered or "full stack" in lowered:
            return "Desenvolvedor Full-Stack"
        if "backend" in lowered:
            return "Desenvolvedor Backend"
        if "frontend" in lowered:
            return "Desenvolvedor Frontend"
        if "software engineer" in lowered:
            return "Software Engineer"
        return None

    def generate_variants(self, user: dict, job: dict, count: int, existing_count: int = 0) -> list[dict]:
        count = max(1, min(count, 10))
        keywords = job.get("requisitos_tecnicos", [])
        educacao = user.get("educacao", [])
        idiomas = user.get("idiomas", [])
        skills = self._prioritize_skills(user.get("habilidades", []), keywords)
        ranked_experiences = self._rank_experiences(user.get("experiencias", []), keywords)

        variants = []
        for offset in range(count):
            round_number = existing_count + offset + 1
            selected_experiences = self._pick_experiences(ranked_experiences, offset)
            generated_content = self._generate_llm_content(
                user=user,
                job=job,
                offset=offset,
                selected_experiences=selected_experiences,
            )
            if generated_content is None:
                generated_content = self._build_rule_based_content(
                    user=user,
                    job=job,
                    offset=offset,
                    selected_experiences=selected_experiences,
                    prioritized_skills=skills,
                )

            generated_content["educacao"] = educacao
            generated_content["idiomas"] = idiomas

            score, reasons = self._score_variant(
                keywords,
                generated_content.get("habilidades", []),
                generated_content.get("experiencias", []),
                generated_content.get("resumo", ""),
            )

            if self.llm_client:
                reasons.insert(0, "Variante gerada com prompt legado focado em ATS.")

            created_at = datetime.now(UTC).isoformat()
            variant_id = uuid.uuid4().hex
            output_path = self.output_dir / f"cv_{variant_id}.docx"
            self._write_docx(generated_content, output_path)

            variants.append(
                {
                    "id": variant_id,
                    "job_id": job["id"],
                    "round": round_number,
                    "ats_score": score,
                    "ats_status": self._status_for_score(score),
                    "ranking_score": score,
                    "motivos": reasons,
                    "content": generated_content,
                    "created_at": created_at,
                    "output_path": str(output_path),
                }
            )

        return variants

    def ensure_docx(self, variant: dict) -> Path:
        output_path = Path(variant["output_path"])
        if not output_path.exists():
            self._write_docx(variant["content"], output_path)
        return output_path

    def _build_llm_client(self) -> Any | None:
        if not self.enable_llm or OpenAI is None or not self.openai_api_key:
            return None
        try:
            return OpenAI(api_key=self.openai_api_key)
        except Exception:
            return None

    def _read_env_value(self, key: str) -> str | None:
        env_value = os.getenv(key)
        if env_value:
            return env_value.strip()
        if not self.env_file.exists():
            return None
        for raw_line in self.env_file.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            current_key, current_value = line.split("=", 1)
            if current_key.strip() == key:
                return current_value.strip().strip('"').strip("'")
        return None

    def _find_first_match(self, lines: list[str], patterns: list[str]) -> str | None:
        for line in lines:
            for pattern in patterns:
                match = re.search(pattern, line, flags=re.IGNORECASE)
                if match:
                    return match.group(1).strip()
        return None

    def extract_keywords(self, content: str) -> list[str]:
        lowered = content.lower()
        found = []
        for keyword in KNOWN_KEYWORDS:
            if keyword in lowered:
                normalized = keyword.replace(".js", "JS").title()
                found.append(normalized)
        if not found:
            tokens = re.findall(r"[A-Za-z][A-Za-z0-9+.#/-]{2,}", content)
            found = [token for token in tokens[:12]]
        return list(dict.fromkeys(found))[:12]

    def extract_behavior_keywords(self, content: str) -> list[str]:
        behaviors = []
        mapping = {
            "lider": "Lideranca",
            "communication": "Comunicacao",
            "comunica": "Comunicacao",
            "ownership": "Ownership",
            "proativo": "Proatividade",
            "proactive": "Proatividade",
            "team": "Trabalho em equipe",
            "colabora": "Colaboracao",
            "analitico": "Perfil analitico",
            "analytical": "Perfil analitico",
        }
        lowered = content.lower()
        for raw, label in mapping.items():
            if raw in lowered:
                behaviors.append(label)
        return list(dict.fromkeys(behaviors))

    def _generate_llm_content(
        self,
        user: dict,
        job: dict,
        offset: int,
        selected_experiences: list[dict],
    ) -> dict | None:
        if not self.llm_client:
            return None

        prompt = self._build_legacy_prompt(user, job, offset, selected_experiences)
        try:
            response = self.llm_client.chat.completions.create(
                model=self.openai_model,
                response_format={"type": "json_object"},
                messages=[
                    {
                        "role": "system",
                        "content": "Voce e um especialista em curriculos ATS e deve responder apenas com JSON valido.",
                    },
                    {"role": "user", "content": prompt},
                ],
            )
            raw_content = response.choices[0].message.content or "{}"
            payload = self._extract_json(raw_content)
            return self._normalize_generated_content(user, job, payload, selected_experiences)
        except Exception:
            return None

    def _build_legacy_prompt(
        self,
        user: dict,
        job: dict,
        offset: int,
        selected_experiences: list[dict],
    ) -> str:
        empresas_validas = [
            {
                "empresa": experience.get("empresa", ""),
                "cargo": experience.get("cargo", ""),
                "periodo": experience.get("periodo", ""),
            }
            for experience in user.get("experiencias", [])
        ]
        project_mapping = self._build_project_mapping(user.get("experiencias", []))
        variation_focus = self._build_variation_focus(job, offset, selected_experiences)

        return f"""
Voce e um especialista em otimizacao de curriculos para ATS.

VAGA ALVO:
- Cargo: {job.get("cargo", "")}
- Empresa: {job.get("empresa", "")}
- ATS: {job.get("ats_detectado", "")}
- Requisitos Tecnicos: {", ".join(job.get("requisitos_tecnicos", []))}
- Requisitos Comportamentais: {", ".join(job.get("requisitos_comportamentais", []))}
- Local: {job.get("local") or "Nao informado"}
- Modalidade: {job.get("modalidade") or "Nao informado"}

DESCRICAO COMPLETA DA VAGA:
{job.get("content", "")}

CURRICULO COMPLETO DO CANDIDATO:
{json.dumps(user, ensure_ascii=False, indent=2)}

EMPRESAS VALIDAS:
{json.dumps(empresas_validas, ensure_ascii=False, indent=2)}

MAPEAMENTO CRITICO DE PROJETOS:
{project_mapping}

CONHECIMENTO DE RH (RAG):
{RAG_CONTEXT}

TAREFA:
Reescreva o curriculo para MAXIMIZAR a chance de passar no ATS da vaga acima.

VARIANTE DESEJADA:
{variation_focus}

REGRA DE OURO (EMPRESAS E PERIODOS):
1. USE APENAS AS EMPRESAS LISTADAS EM "EMPRESAS VALIDAS".
2. SE A EMPRESA DA VAGA NAO ESTIVER NA LISTA DE VALIDAS, NAO A ADICIONE NAS EXPERIENCIAS.
3. NAO MUDE OS CARGOS DRASTICAMENTE.
4. JAMAIS INVENTE UM PERIODO DE TRABALHO.

REGRAS OBRIGATORIAS:
1. Fidelidade a empresas e periodos e obrigatoria.
2. Voce pode e deve mencionar tecnologias da vaga, mesmo quando o curriculo original nao as cita literalmente.
3. Liste no minimo 8 habilidades, preferencialmente 10 a 12, usando apenas tecnologias reais.
4. O resumo deve ser narrativo, forte e focado na vaga. Nao faca uma simples lista seca de tecnologias.
5. Os bullets das experiencias devem ser tecnicos, especificos e plausiveis, evitando genericidade.
6. Se a vaga mencionar arquitetura, design patterns, CI/CD, cloud, observabilidade, qualidade, sustentacao ou integracoes, distribua esses temas nas experiencias de forma natural.
7. Use linguagem de engenheiro de software pleno ou senior, sem exagero artificial.
8. Mantenha o texto natural, mas maximize a cobertura das palavras-chave da vaga.
9. Cada keyword relevante da vaga deve aparecer pelo menos uma vez no resumo ou nas experiencias.
10. Se houver experiencias selecionadas abaixo, priorize-as sem remover coerencia:
{json.dumps(selected_experiences, ensure_ascii=False, indent=2)}
11. Baseie o estilo em resumes de engenharia de software que convertem bem em ATS e com recrutadores: verbos fortes, ownership claro, stack explicita, contexto de sistema e resultado concreto.
12. Evite bullets curtos. Cada bullet deve ter preferencialmente entre 180 e 320 caracteres.
13. Cada experiencia deve ter 4 ou 5 bullets bem desenvolvidos.
14. Sempre que houver evidencias no curriculo base, destaque escalabilidade, integracoes, performance, confiabilidade, reducao de tempo, qualidade de codigo, CI/CD, code review, testes, documentacao e colaboracao com produto ou negocio.
15. Se nao houver metrica numerica confiavel, use impacto operacional realista sem inventar numeros.

RESUMO PROFISSIONAL:
- Estrutura esperada em 2 frases.
- Frase 1: senioridade + area + anos de experiencia + contexto de atuacao.
- Frase 2: expertise em 5 a 7 tecnologias principais + foco tecnico/resultado.

BULLETS DAS EXPERIENCIAS:
- Cada experiencia deve ter 4 a 5 bullets.
- Cada bullet deve citar contexto tecnico, stack, integracao, qualidade, arquitetura, cloud, CI/CD, colaboracao e resultado quando fizer sentido.
- Evite bullets vagos como "Trabalhei com APIs".
- Comece com verbo forte como Desenvolvi, Implementei, Estruturei, Otimizei, Liderei, Arquitetei, Modelei, Automatizei ou Conduzi.
- Mantenha densidade tecnica alta, mas legivel.
- Prefira bullets no estilo:
  "Implementei integracoes REST entre servicos em Java e Node.js, versionando com Git, persistindo em PostgreSQL e documentando contratos para reduzir falhas operacionais."

SAIDA JSON ESPERADA:
{{
  "resumo": "texto",
  "habilidades": ["skill 1", "skill 2"],
  "experiencias": [
    {{
      "empresa": "Nome da empresa original",
      "cargo": "Cargo original",
      "periodo": "Periodo original",
      "descricao": "Descricao curta e coerente",
      "bullets": [
        "Bullet tecnico 1",
        "Bullet tecnico 2",
        "Bullet tecnico 3"
      ]
    }}
  ]
}}

RETORNE APENAS UM JSON VALIDO.
""".strip()

    def _build_project_mapping(self, experiences: list[dict]) -> str:
        lines = []
        for index, experience in enumerate(experiences, start=1):
            bullets = experience.get("realizacoes", [])
            description = experience.get("descricao", "")
            lines.append(
                f"{index}. {experience.get('empresa', '')} | {experience.get('cargo', '')} | {experience.get('periodo', '')}"
            )
            if description:
                lines.append(f"   Contexto base: {description}")
            if bullets:
                for bullet in bullets[:4]:
                    lines.append(f"   Projeto/entrega: {bullet}")
        return "\n".join(lines) if lines else "Sem mapeamento adicional."

    def _build_variation_focus(self, job: dict, offset: int, selected_experiences: list[dict]) -> str:
        focus_labels = [
            "priorize backend, APIs, integracoes e aderencia tecnica direta aos requisitos da vaga",
            "priorize cloud, CI/CD, arquitetura, qualidade e escalabilidade",
            "priorize impacto de produto, colaboracao entre times e experiencia full-stack",
            "priorize sustentacao, operacao, observabilidade e confiabilidade",
        ]
        highlighted = [
            {
                "empresa": experience.get("empresa", ""),
                "cargo": experience.get("cargo", ""),
            }
            for experience in selected_experiences[:2]
        ]
        return (
            f"Variante {offset + 1}: {focus_labels[offset % len(focus_labels)]}. "
            f"Experiencias preferenciais para enfatizar: {json.dumps(highlighted, ensure_ascii=False)}."
        )

    def _extract_json(self, raw_content: str) -> dict[str, Any]:
        content = raw_content.strip()
        fenced_match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", content, flags=re.DOTALL)
        if fenced_match:
            content = fenced_match.group(1).strip()
        if not content.startswith("{"):
            start = content.find("{")
            end = content.rfind("}")
            if start != -1 and end != -1 and end > start:
                content = content[start : end + 1]
        payload = json.loads(content)
        if not isinstance(payload, dict):
            raise ValueError("Invalid JSON payload")
        return payload

    def _normalize_generated_content(
        self,
        user: dict,
        job: dict,
        payload: dict[str, Any],
        selected_experiences: list[dict],
    ) -> dict[str, Any]:
        experiences = self._merge_experiences(selected_experiences, payload.get("experiencias", []))
        experiences = self._harden_experiences_for_ats(
            experiences=experiences,
            base_experiences=selected_experiences,
            job=job,
        )
        skills = self._clean_generated_skills(payload.get("habilidades"), job.get("requisitos_tecnicos", []), user)
        summary = self._sanitize_summary(payload.get("resumo"), user, job, skills, experiences)

        return {
            "nome": user["nome"],
            "cargo": user.get("cargo_atual") or job.get("cargo", ""),
            "email": user["email"],
            "telefone": user["telefone"],
            "linkedin": user["linkedin"],
            "github": user.get("github"),
            "cidade": user["cidade"],
            "estado": user["estado"],
            "resumo": summary,
            "experiencias": experiences,
            "educacao": user.get("educacao", []),
            "habilidades": skills,
            "idiomas": user.get("idiomas", []),
        }

    def _merge_experiences(self, base_experiences: list[dict], generated_experiences: Any) -> list[dict]:
        if not isinstance(generated_experiences, list):
            generated_experiences = []

        merged = []
        for index, base_experience in enumerate(base_experiences):
            generated = generated_experiences[index] if index < len(generated_experiences) else {}
            if not isinstance(generated, dict):
                generated = {}

            bullets = generated.get("bullets") or generated.get("realizacoes") or []
            if isinstance(bullets, str):
                bullets = [item.strip() for item in bullets.split("\n") if item.strip()]
            bullets = [str(item).strip() for item in bullets if str(item).strip()]
            bullets = bullets[:5]

            generated_description = str(generated.get("descricao") or "").strip()
            if not generated_description and bullets:
                generated_description = " ".join(bullets[:2])
            description = generated_description or base_experience.get("descricao", "")

            technologies = self._extract_technologies_from_text(
                " ".join(
                    [
                        description,
                        " ".join(bullets),
                        " ".join(base_experience.get("tecnologias", [])),
                    ]
                )
            )
            if not technologies:
                technologies = base_experience.get("tecnologias", [])

            merged.append(
                {
                    "empresa": base_experience.get("empresa", ""),
                    "cargo": base_experience.get("cargo", ""),
                    "periodo": base_experience.get("periodo", ""),
                    "descricao": description,
                    "tecnologias": technologies[:10],
                    "realizacoes": bullets or base_experience.get("realizacoes", []),
                }
            )

        return merged or base_experiences[:4]

    def _clean_generated_skills(self, raw_skills: Any, keywords: list[str], user: dict) -> list[str]:
        candidates: list[str] = []
        if isinstance(raw_skills, list):
            candidates.extend(str(item).strip() for item in raw_skills if str(item).strip())
        elif isinstance(raw_skills, str):
            candidates.extend(part.strip() for part in raw_skills.split(",") if part.strip())

        candidates.extend(keywords)
        candidates.extend(user.get("habilidades", []))

        cleaned = []
        for skill in candidates:
            normalized = skill.strip()
            if not normalized:
                continue
            lowered = normalized.lower()
            if lowered in INVALID_SKILL_MARKERS:
                continue
            if len(normalized) < 2:
                continue
            if normalized not in cleaned:
                cleaned.append(normalized)
        return cleaned[:12]

    def _sanitize_summary(
        self,
        raw_summary: Any,
        user: dict,
        job: dict,
        skills: list[str],
        experiences: list[dict],
    ) -> str:
        if isinstance(raw_summary, str) and raw_summary.strip():
            return raw_summary.strip()
        return self._build_summary(user, job, skills, user.get("experiencias", []) or experiences)

    def _extract_technologies_from_text(self, text: str) -> list[str]:
        lowered = text.lower()
        found = []
        for keyword in KNOWN_KEYWORDS:
            if keyword in lowered:
                found.append(keyword.replace(".js", "JS").title())
        return list(dict.fromkeys(found))

    def _harden_experiences_for_ats(
        self,
        experiences: list[dict],
        base_experiences: list[dict],
        job: dict,
    ) -> list[dict]:
        hardened = []
        for index, experience in enumerate(experiences):
            base_experience = base_experiences[index] if index < len(base_experiences) else experience
            bullets = self._build_impressive_bullets(experience, base_experience, job)
            technologies = self._combine_technologies(experience, base_experience, job.get("requisitos_tecnicos", []))
            description = self._build_experience_description(experience, base_experience, bullets, technologies)
            hardened.append(
                {
                    "empresa": experience.get("empresa", base_experience.get("empresa", "")),
                    "cargo": experience.get("cargo", base_experience.get("cargo", "")),
                    "periodo": experience.get("periodo", base_experience.get("periodo", "")),
                    "descricao": description,
                    "tecnologias": technologies[:12],
                    "realizacoes": bullets[:5],
                }
            )
        return hardened

    def _build_impressive_bullets(self, experience: dict, base_experience: dict, job: dict) -> list[str]:
        job_keywords = job.get("requisitos_tecnicos", [])
        tech_stack = self._combine_technologies(experience, base_experience, job_keywords)[:6]
        tech_text = ", ".join(tech_stack) if tech_stack else "tecnologias aderentes ao contexto da vaga"
        raw_sources = [
            *(experience.get("realizacoes") or []),
            *(base_experience.get("realizacoes") or []),
        ]
        description_source = experience.get("descricao") or base_experience.get("descricao") or ""
        if description_source:
            raw_sources.append(description_source)
        raw_sources = [str(item).strip().rstrip(".") for item in raw_sources if str(item).strip()]
        raw_sources = list(dict.fromkeys(raw_sources))

        existing_bullets = [
            self._normalize_bullet_text(str(item).strip())
            for item in experience.get("realizacoes", [])
            if str(item).strip()
        ]
        strong_existing = [item for item in existing_bullets if len(item) >= 220]

        if len(strong_existing) >= 4:
            return strong_existing[:5]

        themes = [
            "arquitetura e integracoes",
            "implementacao e persistencia",
            "experiencia do usuario e produto",
            "qualidade, entrega e operacao",
            "colaboracao e lideranca tecnica",
        ]
        bullets = []
        for index, theme in enumerate(themes):
            source = raw_sources[index] if index < len(raw_sources) else description_source
            candidate = self._expand_bullet_for_ats(source, tech_text, job_keywords, theme)
            if candidate and candidate not in bullets:
                bullets.append(candidate)
            if len(bullets) >= min(5, max(4, len(raw_sources))):
                break

        if len(bullets) < 4:
            fallback = self._fallback_bullet(base_experience, tech_text, job_keywords)
            if fallback not in bullets:
                bullets.append(fallback)

        return bullets[:5]

    def _expand_bullet_for_ats(
        self,
        source: str,
        tech_text: str,
        job_keywords: list[str],
        theme: str,
    ) -> str:
        base = self._normalize_bullet_text(self._clean_sentence(source))
        theme_map = {
            "arquitetura e integracoes": (
                "Atuei na estruturacao de fluxos entre servicos e APIs, organizando contratos, integracoes e regras de negocio "
                f"com apoio de {tech_text}, o que aumentou a consistencia da solucao e facilitou a evolucao do produto."
            ),
            "implementacao e persistencia": (
                "Implementei funcionalidades com foco em robustez de backend, modelagem de dados e persistencia, conectando "
                f"camadas aplicacionais e operacionais com {tech_text} para sustentar entregas confiaveis."
            ),
            "experiencia do usuario e produto": (
                "Conectei requisitos de produto e experiencia do usuario com implementacao tecnica, aplicando "
                f"{tech_text} para alinhar backlog, regras de negocio, usabilidade e comportamento da aplicacao com os objetivos de entrega."
            ),
            "qualidade, entrega e operacao": (
                "Apoiei a qualidade tecnica com versionamento, testes, revisao de codigo, documentacao e rotinas de entrega, "
                f"mantendo o ambiente mais previsivel e aderente a praticas modernas em {tech_text}."
            ),
            "colaboracao e lideranca tecnica": (
                "Conduzi alinhamentos tecnicos com times internos, refinando solucoes, analisando impactos e traduzindo necessidades de negocio "
                f"em implementacoes praticas com {tech_text}, mantendo boa comunicacao com stakeholders e clareza nas decisoes tecnicas."
            ),
        }
        complement = theme_map.get(theme, "")
        if not base:
            return complement
        if len(base) < 220:
            return f"{base}. {complement}".strip()
        return base

    def _fallback_bullet(self, base_experience: dict, tech_text: str, job_keywords: list[str]) -> str:
        return (
            "Estruturei entregas com foco em confiabilidade, clareza de implementacao e aderencia tecnica, conectando "
            f"stack, integracoes e rotinas de entrega com {tech_text}, de forma consistente para leitura humana e boa indexacao em ATS."
        )

    def _combine_technologies(self, experience: dict, base_experience: dict, job_keywords: list[str]) -> list[str]:
        combined = []
        for source in [
            experience.get("tecnologias", []),
            base_experience.get("tecnologias", []),
            self._extract_technologies_from_text(experience.get("descricao", "")),
            self._extract_technologies_from_text(" ".join(experience.get("realizacoes", []))),
        ]:
            for item in source:
                normalized = self._normalize_technology_name(str(item).strip())
                if normalized and normalized not in combined:
                    combined.append(normalized)
        return combined

    def _normalize_technology_name(self, value: str) -> str:
        normalized = value.strip()
        mapping = {
            "Node": "Node.js",
            "Nodejs": "Node.js",
            "Node.Js": "Node.js",
            "Javascript": "JavaScript",
            "Typescript": "TypeScript",
            "Ci/Cd": "CI/CD",
            "Api": "API",
            "Apis Rest": "APIs REST",
            "Aws": "AWS",
            "Mysql": "MySQL",
            "Postgresql": "PostgreSQL",
            "Mongodb": "MongoDB",
            "React.Js": "React",
            "Microservices": "Microsserviços",
        }
        return mapping.get(normalized, normalized)

    def _build_experience_description(
        self,
        experience: dict,
        base_experience: dict,
        bullets: list[str],
        technologies: list[str],
    ) -> str:
        description = str(experience.get("descricao") or base_experience.get("descricao") or "").strip()
        if description and len(description) >= 120:
            return description
        company = experience.get("empresa") or base_experience.get("empresa") or "a empresa"
        role = experience.get("cargo") or base_experience.get("cargo") or "a funcao"
        stack = ", ".join(technologies[:5]) if technologies else "stack aderente a vaga"
        opening = self._clean_sentence(description) or f"Atuacao em {company} como {role}"
        return (
            f"{opening}, com foco em entregas de software orientadas a produto e operacao. "
            f"Experiencia reforcada por implementacoes envolvendo {stack}, integracoes entre servicos, qualidade de codigo e colaboracao com areas internas para evolucao continua da solucao."
        )

    def _clean_sentence(self, value: str) -> str:
        cleaned = re.sub(r"\s+", " ", str(value or "").strip())
        cleaned = cleaned.rstrip(".")
        return cleaned

    def _normalize_bullet_text(self, value: str) -> str:
        cleaned = re.sub(r"\s+", " ", value).strip()
        cleaned = re.sub(r"\bDesenvolveu\b", "Desenvolvi", cleaned)
        cleaned = re.sub(r"\bImplementou\b", "Implementei", cleaned)
        cleaned = re.sub(r"\bEstruturou\b", "Estruturei", cleaned)
        cleaned = re.sub(r"\bConduziu\b", "Conduzi", cleaned)
        cleaned = re.sub(r"\bArquitetou\b", "Arquitetei", cleaned)
        cleaned = re.sub(r"\bAutomatizou\b", "Automatizei", cleaned)
        cleaned = re.sub(r"\bAtuou\b", "Atuei", cleaned)
        cleaned = re.sub(r"\bParticipou\b", "Participei", cleaned)
        cleaned = re.sub(r"\bCriou\b", "Criei", cleaned)
        cleaned = re.sub(r"\bLiderou\b", "Liderei", cleaned)
        cleaned = re.sub(r"\bModelou\b", "Modelei", cleaned)
        cleaned = re.sub(r"\bOtimizou\b", "Otimizei", cleaned)
        cleaned = re.sub(r"\bNodejs\b", "Node.js", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bNode\.js\.js\b", "Node.js", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bNode\b(?!\.js)", "Node.js", cleaned)
        cleaned = re.sub(r"\bJavascript\b", "JavaScript", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bApis\b", "APIs", cleaned)
        cleaned = re.sub(r"\bApi\b", "API", cleaned)
        cleaned = re.sub(r"\bAws\b", "AWS", cleaned)
        cleaned = re.sub(r"\bCi/Cd\b", "CI/CD", cleaned)
        cleaned = re.sub(r"\bPostgresql\b", "PostgreSQL", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\bMysql\b", "MySQL", cleaned, flags=re.IGNORECASE)
        return cleaned

    def _build_rule_based_content(
        self,
        user: dict,
        job: dict,
        offset: int,
        selected_experiences: list[dict],
        prioritized_skills: list[str],
    ) -> dict:
        selected_skills = self._rotate_skills(prioritized_skills, offset)
        summary = self._build_summary(user, job, selected_skills, user.get("experiencias", []))
        normalized_experiences = []
        for experience in selected_experiences:
            bullets = experience.get("realizacoes", [])[:5]
            description = experience.get("descricao", "")
            if not description and bullets:
                description = " ".join(bullets[:2])
            normalized_experiences.append(
                {
                    "empresa": experience.get("empresa", ""),
                    "cargo": experience.get("cargo", ""),
                    "periodo": experience.get("periodo", ""),
                    "descricao": description,
                    "tecnologias": experience.get("tecnologias", [])[:10],
                    "realizacoes": bullets,
                }
            )

        normalized_experiences = self._harden_experiences_for_ats(
            experiences=normalized_experiences,
            base_experiences=selected_experiences,
            job=job,
        )

        return {
            "nome": user["nome"],
            "cargo": user.get("cargo_atual") or job.get("cargo", ""),
            "email": user["email"],
            "telefone": user["telefone"],
            "linkedin": user["linkedin"],
            "github": user.get("github"),
            "cidade": user["cidade"],
            "estado": user["estado"],
            "resumo": summary,
            "experiencias": normalized_experiences,
            "educacao": user.get("educacao", []),
            "habilidades": selected_skills,
            "idiomas": user.get("idiomas", []),
        }

    def _rank_experiences(self, experiences: list[dict], keywords: list[str]) -> list[dict]:
        scored = []
        for experience in experiences:
            haystack = " ".join(
                [
                    experience.get("empresa", ""),
                    experience.get("cargo", ""),
                    experience.get("descricao", ""),
                    " ".join(experience.get("tecnologias", [])),
                    " ".join(experience.get("realizacoes", [])),
                ]
            ).lower()
            matches = sum(1 for keyword in keywords if keyword.lower() in haystack)
            scored.append((matches, experience))
        return [experience for _, experience in sorted(scored, key=lambda item: item[0], reverse=True)]

    def _pick_experiences(self, experiences: list[dict], offset: int) -> list[dict]:
        if not experiences:
            return []
        ordered = sorted(experiences, key=self._experience_sort_key, reverse=True)
        return ordered

    def _experience_sort_key(self, experience: dict) -> tuple[int, int]:
        period = str(experience.get("periodo", ""))
        years = [int(value) for value in re.findall(r"(20\d{2})", period)]
        if not years:
            return (0, 0)
        start_year = min(years)
        end_year = max(years)
        lowered = period.lower()
        if "atual" in lowered:
            end_year = datetime.now().year + 1
        return (end_year, start_year)

    def _prioritize_skills(self, skills: list[str], keywords: list[str]) -> list[str]:
        prioritized = []
        remaining = []
        lowered_keywords = [keyword.lower() for keyword in keywords]
        for skill in skills:
            if any(keyword in skill.lower() for keyword in lowered_keywords):
                prioritized.append(skill)
            else:
                remaining.append(skill)
        if not prioritized:
            prioritized = keywords[:6]
        return list(dict.fromkeys(prioritized + remaining))[:16]

    def _rotate_skills(self, skills: list[str], offset: int) -> list[str]:
        if not skills:
            return []
        rotation = offset % max(len(skills), 1)
        rotated = skills[rotation:] + skills[:rotation]
        return rotated[:12]

    def _build_summary(self, user: dict, job: dict, skills: list[str], experiences: list[dict]) -> str:
        top_skills = ", ".join(skills[:7]) if skills else user.get("cargo_atual", "")
        years = self._estimate_years_of_experience(user.get("experiencias", []))
        seniority = "Profissional"
        if years >= 6:
            seniority = "Profissional Senior"
        elif years >= 3:
            seniority = "Profissional Pleno"

        focus = job.get("cargo", user.get("cargo_atual", "profissional de tecnologia"))
        domain = self._infer_domain_from_job(job)
        strengths = self._infer_strengths_from_experiences(experiences)
        return (
            f"{seniority} em {focus} com {years}+ anos de experiencia atuando em {domain}. "
            f"Expertise em {top_skills}, com foco em {strengths} e em gerar entregas tecnicamente solidas, legiveis por ATS e convincentes para avaliacao humana."
        )

    def _estimate_years_of_experience(self, experiences: list[dict]) -> int:
        years = 0
        for experience in experiences:
            period = str(experience.get("periodo", ""))
            matches = [int(value) for value in re.findall(r"(20\d{2})", period)]
            if not matches:
                continue
            start_year = min(matches)
            end_year = max(matches)
            if "atual" in period.lower():
                end_year = max(end_year, datetime.now().year)
            years += max(1, end_year - start_year)
        return max(1, min(12, years))

    def _infer_domain_from_job(self, job: dict) -> str:
        content = f"{job.get('cargo', '')} {job.get('content', '')}".lower()
        if "payment" in content or "pagamento" in content or "pix" in content:
            return "desenvolvimento de solucoes de pagamento e integracoes de alta disponibilidade"
        if "frontend" in content or "react" in content or "angular" in content:
            return "desenvolvimento end-to-end de aplicacoes web com forte interface e integracoes"
        if "backend" in content or "api" in content or "micro" in content:
            return "arquitetura e evolucao de servicos backend, APIs e fluxos de integracao"
        return "desenvolvimento end-to-end de aplicacoes web escalaveis"

    def _infer_strengths_from_experiences(self, experiences: list[dict]) -> str:
        haystack = " ".join(
            [
                experience.get("descricao", "") + " " + " ".join(experience.get("realizacoes", []))
                for experience in experiences
            ]
        ).lower()
        strengths = []
        if any(term in haystack for term in ["api", "integr", "rest", "graphql"]):
            strengths.append("integracoes e desenho de servicos")
        if any(term in haystack for term in ["aws", "docker", "kubernetes", "cloud", "ci/cd"]):
            strengths.append("escalabilidade, entrega continua e operacao")
        if any(term in haystack for term in ["react", "frontend", "ux", "interface"]):
            strengths.append("experiencia de produto e interfaces responsivas")
        if any(term in haystack for term in ["teste", "quality", "qualidade", "review"]):
            strengths.append("qualidade de codigo e manutencao sustentavel")
        if not strengths:
            strengths.append("solucoes de software com foco em impacto operacional")
        return ", ".join(strengths[:3])

    def _score_variant(
        self,
        keywords: list[str],
        skills: list[str],
        experiences: list[dict],
        summary: str,
    ) -> tuple[float, list[str]]:
        reasons = []
        if not keywords:
            return 72.0, ["Vaga sem palavras-chave claras; score estimado por consistencia geral."]

        keyword_hits = sum(1 for keyword in keywords if keyword.lower() in " ".join(skills).lower())
        experience_hits = 0
        for keyword in keywords:
            needle = keyword.lower()
            if any(
                needle in " ".join(
                    [
                        experience.get("cargo", ""),
                        experience.get("descricao", ""),
                        " ".join(experience.get("tecnologias", [])),
                        " ".join(experience.get("realizacoes", [])),
                    ]
                ).lower()
                for experience in experiences
            ):
                experience_hits += 1

        summary_hits = sum(1 for keyword in keywords if keyword.lower() in summary.lower())
        max_hits = max(len(keywords), 1)

        keyword_score = keyword_hits / max_hits
        experience_score = experience_hits / max_hits
        summary_score = summary_hits / max_hits

        raw_score = (keyword_score * 0.45) + (experience_score * 0.4) + (summary_score * 0.15)
        score = round(min(98.0, 58 + (raw_score * 40)), 1)

        if keyword_hits:
            reasons.append(f"{keyword_hits} requisitos tecnicos foram refletidos nas habilidades.")
        if experience_hits:
            reasons.append(f"{experience_hits} requisitos aparecem nas experiencias selecionadas.")
        if summary_hits:
            reasons.append("O resumo foi alinhado ao foco da vaga.")
        if not reasons:
            reasons.append("CV estruturado a partir do perfil ativo, com baixa aderencia detectada aos termos da vaga.")

        return score, reasons

    def _status_for_score(self, score: float) -> str:
        if score >= 85:
            return "APPROVED"
        if score >= 70:
            return "RISK"
        return "REJECTED"

    def _write_docx(self, content: dict, output_path: Path) -> None:
        document = Document()
        self._configure_document(document)

        name = document.add_paragraph()
        name.paragraph_format.space_after = Pt(1.5)
        name_run = name.add_run(content["nome"])
        self._style_run(name_run, size=Pt(15.96), bold=True)

        subtitle = document.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(4)
        subtitle_run = subtitle.add_run(content.get("cargo", ""))
        self._style_run(subtitle_run, size=Pt(12), bold=True, font_name=FONT_ITALIC)

        location_parts = [part for part in [content.get("cidade"), content.get("estado")] if part]
        if location_parts:
            location = document.add_paragraph()
            location.paragraph_format.space_after = Pt(1)
            location_text = location_parts[0] if len(location_parts) == 1 else f"{location_parts[0]} \u2013 {location_parts[1]}"
            location_run = location.add_run(location_text)
            self._style_run(location_run, size=Pt(11.04), font_name="Arial")

        email = content.get("email")
        if email:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(email)
            self._style_run(run, size=Pt(11.04), font_name="Arial")

        phone = content.get("telefone")
        if phone:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(f"Telefone: {phone}")
            self._style_run(run, size=Pt(11.04), font_name=FONT_ITALIC)

        for link in [content.get("linkedin"), content.get("github")]:
            if link:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(1)
                self._add_hyperlink(paragraph, link, link)

        self._add_section_heading(document, "RESUMO")
        summary = document.add_paragraph()
        summary.paragraph_format.space_after = Pt(9)
        summary.paragraph_format.line_spacing = 1.1
        summary_run = summary.add_run(content.get("resumo", ""))
        self._style_run(summary_run, size=Pt(11.04), font_name=FONT_ITALIC)

        if content.get("habilidades"):
            self._add_section_heading(document, "HABILIDADES")
            for skill in content["habilidades"]:
                skill_paragraph = document.add_paragraph()
                skill_paragraph.paragraph_format.space_after = Pt(0.8)
                skill_run = skill_paragraph.add_run(skill)
                self._style_run(skill_run, size=Pt(11.04), font_name=FONT_ITALIC)

        if content.get("educacao"):
            self._add_section_heading(document, "FORMA\u00c7\u00c3O")
            for education in content["educacao"]:
                course = document.add_paragraph()
                course.paragraph_format.space_after = Pt(0.8)
                course_run = course.add_run(f"{education.get('curso', '')} \u2013 {education.get('instituicao', '')}")
                self._style_run(course_run, size=Pt(11.04), font_name=FONT_ITALIC)

                period = document.add_paragraph()
                period.paragraph_format.space_after = Pt(6)
                period_run = period.add_run(education.get("periodo", ""))
                self._style_run(period_run, size=Pt(11.04), font_name=FONT_ITALIC)

        if content.get("experiencias"):
            self._add_section_heading(document, "EXPERI\u00caNCIA PROFISSIONAL", draw_line=False)

        for experience in content.get("experiencias", []):
            company = document.add_paragraph()
            company.paragraph_format.space_after = Pt(0.8)
            company_run = company.add_run(experience.get("empresa", ""))
            self._style_run(company_run, size=Pt(9.96), bold=True)

            role = document.add_paragraph()
            role.paragraph_format.space_after = Pt(0.5)
            role_run = role.add_run(experience.get("cargo", ""))
            self._style_run(role_run, size=Pt(11.04), bold=True, italic=True, font_name=FONT_ITALIC)

            if experience.get("periodo"):
                period = document.add_paragraph()
                period.paragraph_format.space_after = Pt(2.5)
                period_run = period.add_run(experience["periodo"])
                self._style_run(period_run, size=Pt(11.04), bold=True, italic=True, font_name=FONT_ITALIC)

            for item in experience.get("realizacoes", []):
                bullet = document.add_paragraph()
                bullet.paragraph_format.left_indent = Mm(4.5)
                bullet.paragraph_format.first_line_indent = Mm(-2)
                bullet.paragraph_format.space_after = Pt(1.5)
                bullet.paragraph_format.line_spacing = 1.08
                bullet_run = bullet.add_run(f"\u2022  {item}")
                self._style_run(bullet_run, size=Pt(11.04), font_name=FONT_ITALIC)

            if experience.get("descricao") and not experience.get("realizacoes"):
                description = document.add_paragraph()
                description.paragraph_format.space_after = Pt(4)
                description.paragraph_format.line_spacing = 1.08
                description_run = description.add_run(experience["descricao"])
                self._style_run(description_run, size=Pt(11.04), font_name=FONT_ITALIC)

        if content.get("idiomas"):
            self._add_section_heading(document, "IDIOMAS")
            for language in content["idiomas"]:
                language_paragraph = document.add_paragraph()
                language_paragraph.paragraph_format.space_after = Pt(1)
                language_run = language_paragraph.add_run(f"{language.get('idioma', '')}: {language.get('nivel', '')}")
                self._style_run(language_run, size=Pt(11.04), font_name=FONT_ITALIC)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15.5)
        section.bottom_margin = Mm(16)
        section.left_margin = Mm(17)
        section.right_margin = Mm(17)

        normal_style = document.styles["Normal"]
        normal_style.font.name = FONT_ITALIC
        normal_style.font.size = Pt(11.04)
        normal_style.font.color.rgb = TEXT_COLOR

    def _add_section_heading(self, document: Document, title: str, draw_line: bool = True) -> None:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(7)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(title)
        self._style_run(run, size=Pt(12.96), bold=True)
        if draw_line:
            border = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "A0A0A0")
            border.append(bottom)
            heading._p.get_or_add_pPr().append(border)

    def _style_run(
        self,
        run,
        *,
        size: Pt,
        bold: bool = False,
        italic: bool = False,
        font_name: str = FONT_REGULAR,
        color: RGBColor = TEXT_COLOR,
    ) -> None:
        run.bold = bold
        run.italic = italic
        run.font.name = font_name
        run.font.size = size
        run.font.color.rgb = color

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        part = paragraph.part
        relationship_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "366171")
        run_properties.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(underline)

        run_fonts = OxmlElement("w:rFonts")
        run_fonts.set(qn("w:ascii"), "Arial")
        run_fonts.set(qn("w:hAnsi"), "Arial")
        run_properties.append(run_fonts)

        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "22")
        run_properties.append(size)

        run.append(run_properties)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)
