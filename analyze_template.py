"""
Script to analyze CV template structure
"""
from docx import Document
from pathlib import Path

template_path = "layoutCV/layout.docx"
doc = Document(template_path)

print("="*60)
print("CV TEMPLATE ANALYSIS")
print("="*60)
print(f"\nFile: {template_path}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Sections: {len(doc.sections)}")
print(f"Tables: {len(doc.tables)}")

# Analyze sections
print("\n" + "="*60)
print("SECTIONS")
print("="*60)
for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  Page width: {section.page_width.inches:.2f} inches")
    print(f"  Page height: {section.page_height.inches:.2f} inches")
    print(f"  Left margin: {section.left_margin.inches:.2f} inches")
    print(f"  Right margin: {section.right_margin.inches:.2f} inches")

# Analyze paragraphs
print("\n" + "="*60)
print("PARAGRAPHS (first 30)")
print("="*60)
for i, p in enumerate(doc.paragraphs[:30]):
    style = p.style.name if p.style else "None"
    text = p.text.strip()[:60]
    if text:
        print(f"{i:3d} | {style:20s} | {text}")

# Check for placeholders
print("\n" + "="*60)
print("PLACEHOLDERS DETECTED")
print("="*60)
import re
placeholders = set()
for p in doc.paragraphs:
    matches = re.findall(r'\{\{([A-Z_]+)\}\}', p.text)
    placeholders.update(matches)

if placeholders:
    print(f"Found {len(placeholders)} placeholders:")
    for ph in sorted(placeholders):
        print(f"  - {{{{{ph}}}}}")
else:
    print("No {{PLACEHOLDER}} markers found - using section-based approach")

# Analyze styles
print("\n" + "="*60)
print("STYLES USED")
print("="*60)
styles_used = {}
for p in doc.paragraphs:
    if p.style:
        style_name = p.style.name
        styles_used[style_name] = styles_used.get(style_name, 0) + 1

for style, count in sorted(styles_used.items()):
    print(f"  {style}: {count}x")

print("\n" + "="*60)
